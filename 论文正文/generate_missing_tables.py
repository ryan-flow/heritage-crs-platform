# -*- coding: utf-8 -*-
"""
生成第15、16张数据表的Word文档
格式完全按照示例：
（X）表名 表结构如下表X所示：
表X 表名
字段名称	类型	长度	字段说明	主键
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

FONT_NAME_CN = "黑体"
FONT_SIZE = Pt(10.5)  # 五号字


def format_cell(cell):
    """格式化单元格：字体五号黑体，不加粗，居中，行间距1.5倍"""
    for para in cell.paragraphs:
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in para.runs:
            run.font.name = FONT_NAME_CN
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
            run.font.size = FONT_SIZE
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.bold = False


def add_table_title(doc, num, table_name_cn, table_name_en):
    """添加表标题"""
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    run = para.add_run(f"（{num}）{table_name_cn}（{table_name_en}）表结构如下表{num}所示：")
    run.font.name = FONT_NAME_CN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
    run.font.size = FONT_SIZE
    run.font.bold = False


def add_table_label(doc, num, table_name_cn):
    """添加表X 表名"""
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    run = para.add_run(f"表{num} {table_name_cn}")
    run.font.name = FONT_NAME_CN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
    run.font.size = FONT_SIZE
    run.font.bold = False


def create_table(doc, fields):
    """创建表格"""
    table = doc.add_table(rows=len(fields)+1, cols=5)
    table.style = 'Table Grid'
    
    headers = ["字段名称", "类型", "长度", "字段说明", "主键"]
    for col, header in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = header
    
    for row_idx, field in enumerate(fields):
        for col_idx, value in enumerate(field):
            cell = table.rows[row_idx+1].cells[col_idx]
            cell.text = str(value)
    
    for row in table.rows:
        for cell in row.cells:
            format_cell(cell)


def add_empty_lines(doc, count=2):
    """添加空行"""
    for _ in range(count):
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = 1.5


def main():
    doc = Document()
    
    # 第15张表：活动报名表
    activity_registration_fields = [
        ["id", "bigint", "", "主键", "主键"],
        ["activity_id", "bigint", "", "活动编号", ""],
        ["user_id", "bigint", "", "用户编号", ""],
        ["remark", "varchar", "255", "报名备注", ""],
        ["status", "varchar", "20", "报名状态", ""],
        ["created_at", "timestamp", "", "创建时间", ""],
        ["updated_at", "timestamp", "", "更新时间", ""]
    ]
    add_table_title(doc, 15, "活动报名表", "activity_registrations")
    add_table_label(doc, 15, "活动报名表")
    create_table(doc, activity_registration_fields)
    add_empty_lines(doc, 2)
    
    # 第16张表：讨论话题表
    discussion_topics_fields = [
        ["id", "bigint", "", "主键", "主键"],
        ["user_id", "bigint", "", "发帖用户编号", ""],
        ["nickname", "varchar", "64", "发帖人昵称", ""],
        ["title", "varchar", "200", "话题标题", ""],
        ["content", "longtext", "4294967295", "话题内容", ""],
        ["cover_url", "varchar", "255", "封面图地址", ""],
        ["image_urls", "text", "", "图片地址列表(JSON)", ""],
        ["is_featured", "boolean", "", "是否精选", ""],
        ["like_count", "int", "", "点赞数", ""],
        ["favorite_count", "int", "", "收藏数", ""],
        ["comment_count", "int", "", "评论数", ""],
        ["created_at", "timestamp", "", "创建时间", ""]
    ]
    add_table_title(doc, 16, "讨论话题表", "discussion_topics")
    add_table_label(doc, 16, "讨论话题表")
    create_table(doc, discussion_topics_fields)
    
    output_path = "d:\\桌面\\毕业设计\\论文正文\\缺失数据表.docx"
    doc.save(output_path)
    
    print(f"生成完成！")
    print(f"- 包含表15: 活动报名表（activity_registrations）")
    print(f"- 包含表16: 讨论话题表（discussion_topics）")
    print(f"- 格式: 五号黑体，行间距1.5倍，居中对齐")
    print(f"- 输出文件: {output_path}")


if __name__ == "__main__":
    main()
