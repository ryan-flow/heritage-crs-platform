# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

path = r"d:\桌面\毕业设计\论文正文\最终正式版\大数据222-202210274225-王子轩-基于CRS推荐与AI数字人的非遗文化传播系统实现.docx"

try:
    doc = Document(path)
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else "Normal"
            print(f"[P{i}|{style}] {text[:200]}")
    
    print(f"\n===== 总段落数: {len(doc.paragraphs)} =====")
    print(f"===== 总表格数: {len(doc.tables)} =====")
    
    for ti, table in enumerate(doc.tables):
        print(f"\n--- 表格{ti+1} ({len(table.rows)}行 x {len(table.columns)}列) ---")
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip()[:50] for cell in row.cells]
            print(f"  R{ri}: {' | '.join(cells)}")
            if ri > 3:
                print(f"  ... (省略{len(table.rows)-4}行)")
                break
except Exception as e:
    print(f"错误: {e}")
    print(f"Python版本: {sys.version}")
    print(f"尝试安装: pip install python-docx")
