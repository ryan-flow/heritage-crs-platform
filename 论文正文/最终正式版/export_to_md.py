# -*- coding: utf-8 -*-
"""导出docx到markdown，保留结构和表格"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

docx_path = r"d:\桌面\毕业设计\论文正文\最终正式版\大数据222-202210274225-王子轩-基于CRS推荐与AI数字人的非遗文化传播系统实现.docx"
md_path = r"d:\桌面\毕业设计\论文正文\最终正式版\论文全文.md"

doc = Document(docx_path)
lines = []

def get_style(para):
    style = para.style.name if para.style else "Normal"
    return style

for i, para in enumerate(doc.paragraphs):
    text = para.text
    if not text:
        continue
    style = get_style(para)
    
    # 判断是否为标题
    if "Heading" in style or "heading" in style or "标题" in style:
        level = 1
        for key in ["Heading 1", "heading 1", "TOCHeading"]:
            if key in style:
                level = 1
                break
        if "Heading 2" in style or "heading 2" in style:
            level = 2
        elif "Heading 3" in style or "heading 3" in style:
            level = 3
        elif "Heading 4" in style or "heading 4" in style:
            level = 4
        
        prefix = "#" * level
        lines.append(f"\n{prefix} {text}\n")
    else:
        lines.append(f"\n{text}\n")

# 处理表格
for ti, table in enumerate(doc.tables):
    lines.append(f"\n**表格{ti+1}**\n")
    lines.append("\n| " + " | ".join([cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells]) + " |")
    lines.append("|" + "|".join(["---" for _ in table.rows[0].cells]) + "|")
    for row in table.rows[1:]:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        lines.append("| " + " | ".join(cells) + " |")

with open(md_path, 'w', encoding='utf-8') as f:
    f.write("".join(lines))

print(f"导出完成！")
print(f"段落数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")
print(f"输出文件: {md_path}")
print(f"文件大小: {os.path.getsize(md_path)} 字节")
