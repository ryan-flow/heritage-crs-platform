# -*- coding: utf-8 -*-
"""
批量处理docx表格格式（5.4第二版）：
1. 不修改重复标题设置（保持原样）
2. 确保所有单元格行间距严格1.5倍
3. 确保所有字体五号黑体不加粗（包括表头）
4. 确保所有单元格格式一致
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

FONT_NAME_CN = "黑体"
FONT_SIZE = Pt(10.5)  # 五号字


def format_cell(cell):
    """格式化单元格：字体五号黑体，不加粗，居中，行间距1.5倍"""
    for para in cell.paragraphs:
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in para.runs:
            run.font.name = FONT_NAME_CN
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
            run.font.size = FONT_SIZE
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.bold = False  # 不加粗


def format_all_tables_in_doc(input_path, output_path):
    """处理文档中所有表格"""
    doc = Document(input_path)

    table_count = 0

    for i, table in enumerate(doc.tables):
        table_count += 1
        
        for row in table.rows:
            for cell in row.cells:
                format_cell(cell)

    doc.save(output_path)

    print(f"处理完成！")
    print(f"- 总表格数: {table_count}")
    print(f"- 字体: 五号黑体（不加粗，包括表头）")
    print(f"- 行距: 严格1.5倍")
    print(f"- 对齐: 居中")
    print(f"- 输出文件: {output_path}")


if __name__ == "__main__":
    input_file = "d:\\桌面\\毕业设计\\论文正文\\5.3第一版.docx"
    output_file = "d:\\桌面\\毕业设计\\论文正文\\5.4第二版.docx"

    format_all_tables_in_doc(input_file, output_file)
