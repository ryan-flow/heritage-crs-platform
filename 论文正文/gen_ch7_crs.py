# -*- coding: utf-8 -*-
"""
生成修正后的7.4.1和7.4.2节Word文档
基于实际代码置信度计算公式:
  S_explicit = heritage*30 + scene*20 + region*20 + ask_answers*10  (上限100)
  S_implicit = browse*5 + registration*15 + engagement*8
               + (ask激励: ask_count>0 ? ask_count*7 : 0)
               + (qa激励: qa_count>0 ? qa_count*5 : 0)
  S_dialogue = qa_count*4 + ask_count*12  (上限100)
  C = 0.40*Se + 0.35*Si + 0.25*Sd

阈值: C<28 cold_start / 28<=C<62 mixed / C>=62 precision
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def calc_implicit(browse, registration, engagement, ask_count, qa_count):
    si = browse * 5 + registration * 15 + engagement * 8
    if ask_count > 0:
        si += min(20, ask_count * 7)
    if qa_count > 0:
        si += min(20, qa_count * 5)
    return min(100, si)

def calc_confidence(se, si, sd):
    return round(0.40 * se + 0.35 * si + 0.25 * sd, 2)

def determine_mode(c):
    if c < 28:
        return "cold_start"
    elif c < 62:
        return "mixed"
    else:
        return "precision"

# ===== 有CRS引导场景(表34) =====
# 交互设计: 8轮CRS引导对话
# 轮次  动作           ask回答  browse  registration  engagement  qa提问
# R1    询问类别→回答→浏览   1       1       0           0          1
# R2    询问地域→回答       2       1       0           0          2
# R3    用户浏览内容        2       2       0           0          2
# R4    询问场景→回答       3       2       0           0          3
# R5    用户浏览内容        3       3       0           0          3
# R6    询问深度→回答       4       3       0           0          4
# R7    点击推荐内容        4       4       0           1          4
# R8    活动报名            4       4       1           1          4

table34_data = []
for r in range(1, 9):
    if r == 1:
        ask_count, browse, registration, engagement, qa_count = 1, 1, 0, 0, 1
        action = "询问类别"
    elif r == 2:
        ask_count, browse, registration, engagement, qa_count = 2, 1, 0, 0, 2
        action = "询问地域"
    elif r == 3:
        ask_count, browse, registration, engagement, qa_count = 2, 2, 0, 0, 2
        action = "浏览内容"
    elif r == 4:
        ask_count, browse, registration, engagement, qa_count = 3, 2, 0, 0, 3
        action = "询问场景"
    elif r == 5:
        ask_count, browse, registration, engagement, qa_count = 3, 3, 0, 0, 3
        action = "浏览内容"
    elif r == 6:
        ask_count, browse, registration, engagement, qa_count = 4, 3, 0, 0, 4
        action = "询问深度"
    elif r == 7:
        ask_count, browse, registration, engagement, qa_count = 4, 4, 0, 1, 4
        action = "点击推荐"
    else:
        ask_count, browse, registration, engagement, qa_count = 4, 4, 1, 1, 4
        action = "活动报名"

    se = min(100, ask_count * 10)
    si = calc_implicit(browse, registration, engagement, ask_count, qa_count)
    sd = min(100, qa_count * 4 + ask_count * 12)
    c = calc_confidence(se, si, sd)
    mode = determine_mode(c)
    table34_data.append([f"第{r}轮", action, se, si, sd, c, mode])
    print(f"第{r}轮 {action}: Se={se} Si={si} Sd={sd} C={c} mode={mode}")

print()

# ===== 无CRS引导场景(表35) =====
# 仅被动浏览+点击，无ASK回答，无显式偏好，无对话信号
# browse随轮次递增，有点击时engagement+1

table35_data = []
browse = 0
engagement = 0
for r in range(1, 9):
    if r == 1:
        browse = 1; registration = 0; engagement = 0; ask_count = 0; qa_count = 0
        action = "用户浏览"
    elif r == 2:
        browse = 2; registration = 0; engagement = 0; ask_count = 0; qa_count = 0
        action = "用户浏览"
    elif r == 3:
        browse = 3; registration = 0; engagement = 0; ask_count = 0; qa_count = 0
        action = "用户浏览"
    elif r == 4:
        browse = 3; registration = 0; engagement = 1; ask_count = 0; qa_count = 0
        action = "用户点击"
    elif r == 5:
        browse = 4; registration = 0; engagement = 1; ask_count = 0; qa_count = 0
        action = "用户浏览"
    elif r == 6:
        browse = 5; registration = 0; engagement = 1; ask_count = 0; qa_count = 0
        action = "用户浏览"
    elif r == 7:
        browse = 5; registration = 0; engagement = 2; ask_count = 0; qa_count = 0
        action = "用户点击"
    else:
        browse = 6; registration = 0; engagement = 2; ask_count = 0; qa_count = 0
        action = "用户浏览"

    se = 0
    si = calc_implicit(browse, registration, engagement, ask_count, qa_count)
    sd = 0
    c = calc_confidence(se, si, sd)
    mode = determine_mode(c)
    table35_data.append([f"第{r}轮", action, se, si, sd, c, mode])
    print(f"第{r}轮 {action}: Se={se} Si={si} Sd={sd} C={c} mode={mode}")

print(f"\n有CRS: 最终C={table34_data[-1][5]}, 无CRS: 最终C={table35_data[-1][5]}")
ratio = round(table34_data[-1][5] / table35_data[-1][5], 2)
print(f"倍数: {ratio}")

# ===== 生成Word =====
doc = Document()

style = doc.styles["Normal"]
style.font.name = "宋体"
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "黑体"
    hs.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    hs.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        hs.font.size = Pt(16)
        hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        hs.font.size = Pt(14)
    else:
        hs.font.size = Pt(13)

doc.add_paragraph()
doc.add_heading("7.4.1  实验设计", level=2)
p1 = doc.add_paragraph()
p1.add_run("本节关注CRS机制对冷启动问题的改善作用，比较有CRS引导和无CRS引导两种情况下新用户置信度的增长过程，以及达到精准推荐模式所需的交互轮次。实验采用三维置信度模型来模拟新用户从注册到获得有效推荐的完整过程。三维置信度评估当前对用户的理解程度，包括显式偏好维度、隐式行为维度和对话维度。设显式偏好得分为Se，隐式行为得分为Si，对话得分为Sd，综合置信度可以表示为：")
p_formula = doc.add_paragraph("C = 0.40Se + 0.35Si + 0.25Sd")
p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p_formula.runs:
    run.bold = True
    run.font.size = Pt(12)
p2 = doc.add_paragraph()
p2.add_run("当C<28时为冷启动模式，28≤C<62时为混合模式，C≥62时为精准模式。有CRS引导的情况下，系统通过ASK模板主动询问用户的类别偏好、地域偏好、场景偏好和了解程度，同时用户也会产生浏览、点击等隐式行为；无CRS引导的情况下，用户只能通过被动浏览积累隐式行为信号，没有显式偏好输入和对话信号。")

doc.add_paragraph()
doc.add_heading("7.4.2  实验结果", level=2)

p3 = doc.add_paragraph()
r = p3.add_run("表34  CRS机制下冷启动置信度变化过程")
r.bold = True

table34 = doc.add_table(rows=1, cols=7)
table34.style = "Table Grid"
hdr_cells = table34.rows[0].cells
headers = ["交互步骤", "动作", "S\u660e\u786e\u5047\u597d", "S\u9690\u5f0f\u884c\u4e3a", "S\u5bf9\u8bdd\u8bed\u4e49", "\u7f6e\u4fe1\u5ea6C", "CRS\u6a21\u5f0f"]
for i, (cell, h) in enumerate(zip(hdr_cells, headers)):
    cell.text = h
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(10)
    set_cell_bg(cell, "D9E2F3")

for row_data in table34_data:
    row = table34.add_row()
    for i, (cell, val) in enumerate(zip(row.cells, row_data)):
        cell.text = str(val)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)

doc.add_paragraph()
p4 = doc.add_paragraph()
r = p4.add_run("表35  无CRS机制下冷启动置信度变化过程")
r.bold = True

table35 = doc.add_table(rows=1, cols=7)
table35.style = "Table Grid"
hdr35 = table35.rows[0].cells
for i, (cell, h) in enumerate(zip(hdr35, headers)):
    cell.text = h
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(10)
    set_cell_bg(cell, "D9E2F3")

for row_data in table35_data:
    row = table35.add_row()
    for i, (cell, val) in enumerate(zip(row.cells, row_data)):
        cell.text = str(val)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)

doc.add_paragraph()
doc.add_heading("7.4.3  结果分析", level=2)
p5 = doc.add_paragraph()
c_crs = table34_data[-1][5]
c_no_crs = table35_data[-1][5]
ratio = round(c_crs / c_no_crs, 2)
se_last = table34_data[-1][2]
sd_last = table34_data[-1][4]
weight = round(0.40 + 0.25, 2)
p5.add_run(f"从实验结果可以看出，CRS机制对冷启动问题的改善效果非常明显。以置信度增长速度为观察角度，有CRS引导的用户在第8轮交互时置信度达到{c_crs}，进入混合模式，而无CRS引导的用户经过8轮交互后置信度仅为{c_no_crs}，始终停留在冷启动模式。以最终置信度为对比，CRS引导下的最终置信度{c_crs}是无CRS的{c_no_crs}的{ratio}倍，差距悬殊。以信号来源为分析维度，CRS引导下显式偏好分数S_explicit从0增长到{se_last}，对话分数S_dialogue从0增长到{sd_last}，这两个维度是无CRS情况下完全无法获得的，而它们在置信度计算中的权重合计达到{weight}，是决定性的信号来源。这说明CRS通过主动提问来获取用户的显式偏好和对话信号，确实是解决冷启动问题最有效的技术路径。")

out_path = "d:\\桌面\\毕业设计\\论文正文\\ch7_4_crs_experiment_fixed.docx"
doc.save(out_path)
print(f"\n已保存: {out_path}")