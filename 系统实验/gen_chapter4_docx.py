#!/usr/bin/env python3
"""生成第4章系统设计Word文档（三线表格式）"""

import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

MD_PATH = r"d:\桌面\毕业设计\论文正文\第四章-系统设计.md"
OUT_PATH = r"d:\桌面\毕业设计\论文正文\第四章-系统设计.docx"


def _set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _make_three_line_table(doc, headers, rows, caption=""):
    if caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap_para.add_run(caption)
        run.font.size = Pt(10.5)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        cap_para.paragraph_format.space_after = Pt(4)
        cap_para.paragraph_format.space_before = Pt(6)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(header))
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        _set_cell_shading(cell, "F2F2F2")

    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()
            tc_borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'</w:tcBorders>'
            )
            existing = tc_pr.find(qn("w:tcBorders"))
            if existing is not None:
                tc_pr.remove(existing)
            tc_pr.append(tc_borders)

    for cell in table.rows[0].cells:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
            f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tcBorders>'
        )
        existing = tc_pr.find(qn("w:tcBorders"))
        if existing is not None:
            tc_pr.remove(existing)
        tc_pr.append(tc_borders)

    for cell in table.rows[-1].cells:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tcBorders>'
        )
        existing = tc_pr.find(qn("w:tcBorders"))
        if existing is not None:
            tc_pr.remove(existing)
        tc_pr.append(tc_borders)

    return table


def _add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    return heading


def _add_body_text(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return para


def _add_formula(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.font.italic = True
    return para


def main():
    with open(MD_PATH, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing = 1.5

    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        if line.startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1
            continue

        if line.startswith("（图"):
            _add_body_text(doc, line)
            i += 1
            continue

        if line.startswith("C = 0.40"):
            _add_formula(doc, line)
            i += 1
            continue

        if line.startswith("#") or re.match(r"^4[\.\s]", line):
            level = 0
            if line.startswith("4 系统设计") and not line.startswith("4."):
                level = 1
            elif re.match(r"^4\.\d+ ", line):
                sub = line.split()[0]
                dots = sub.count(".")
                level = dots + 1
            if level > 0:
                _add_heading(doc, line, level=min(level, 4))
                i += 1
                continue

        if line.startswith("|") and "|" in line[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1

            if len(table_lines) < 2:
                continue

            caption = ""
            for j in range(i - len(table_lines) - 1, -1, -1):
                prev = lines[j].strip()
                if prev.startswith("表") and ("4-" in prev or "表4" in prev):
                    caption = prev
                    break
                if prev and not prev.startswith("|") and not prev.startswith("-"):
                    break

            headers = [c.strip() for c in table_lines[0].split("|")[1:-1]]
            rows = []
            for tl in table_lines[2:]:
                cells = [c.strip() for c in tl.split("|")[1:-1]]
                if len(cells) == len(headers):
                    rows.append(cells)

            _make_three_line_table(doc, headers, rows, caption=caption)
            continue

        _add_body_text(doc, line)
        i += 1

    doc.save(OUT_PATH)
    print(f"Word文档已生成: {OUT_PATH}")


if __name__ == "__main__":
    main()
