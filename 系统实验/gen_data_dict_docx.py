#!/usr/bin/env python3
"""生成3.5系统数据字典Word文档（三线表格式）"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

HEADERS = ["字段名称", "类型", "长度", "字段说明", "主键", "默认值"]

TABLES = [
    {
        "num": 1,
        "name": "用户表（users）",
        "desc": "储存系统用户的基本信息、角色标识和CRS推荐画像数据",
        "fields": [
            ["id", "bigint", "", "主键", "主键", ""],
            ["openid", "varchar", "64", "微信唯一标识", "", ""],
            ["nickname", "varchar", "64", "用户昵称", "", ""],
            ["phone", "varchar", "20", "手机号", "", ""],
            ["avatar_url", "varchar", "255", "头像地址", "", ""],
            ["role", "varchar", "20", "角色", "", "user"],
            ["is_active", "boolean", "", "是否活跃", "", "1"],
            ["preferred_heritage_types", "text", "", "非遗类型偏好(JSON数组)", "", ""],
            ["preferred_scene_types", "text", "", "场景偏好(JSON数组)", "", ""],
            ["preferred_regions", "text", "", "地区偏好(JSON数组)", "", ""],
            ["confidence_score", "real", "", "CRS综合置信度(0-100)", "", "0"],
            ["score_explicit", "real", "", "显式偏好维度分(0-100)", "", "0"],
            ["score_implicit", "real", "", "隐式行为维度分(0-100)", "", "0"],
            ["score_dialogue", "real", "", "对话语义维度分(0-100)", "", "0"],
            ["created_at", "timestamp", "", "创建时间", "", "CURRENT_TIMESTAMP"],
        ],
    },
    {
        "num": 2,
        "name": "内容表（contents）",
        "desc": "储存非遗内容的核心信息，涵盖标题、正文、来源、审核状态和推荐评分",
        "fields": [
            ["id", "bigint", "", "主键", "主键", ""],
            ["title", "varchar", "200", "标题", "", ""],
            ["cover_url", "varchar", "255", "封面图地址", "", ""],
            ["content_type", "varchar", "30", "内容类型", "", ""],
            ["source_site", "varchar", "120", "来源站点", "", ""],
            ["source_url", "varchar", "500", "来源链接", "", ""],
            ["summary", "text", "", "摘要", "", ""],
            ["body", "longtext", "4294967295", "正文", "", ""],
            ["chapter", "varchar", "120", "所属章节", "", ""],
            ["sub_chapter", "varchar", "120", "子章节", "", ""],
            ["content_hash", "varchar", "40", "正文哈希(去重)", "", ""],
            ["quality_score", "float", "", "质量评分", "", "0"],
            ["review_status", "varchar", "20", "审核状态", "", "pending"],
            ["import_batch", "varchar", "60", "导入批次", "", ""],
            ["is_featured", "boolean", "", "是否精选", "", "0"],
            ["status", "varchar", "20", "发布状态", "", "draft"],
            ["published_at", "timestamp", "", "发布时间", "", ""],
            ["created_by", "bigint", "", "创建人编号", "", ""],
            ["created_at", "timestamp", "", "创建时间", "", "CURRENT_TIMESTAMP"],
        ],
    },
    {
        "num": 3,
        "name": "内容收藏表（content_favorites）",
        "desc": "储存用户对非遗内容的收藏记录",
        "fields": [
            ["id", "bigint", "", "主键", "主键", ""],
            ["user_id", "bigint", "", "用户编号", "", ""],
            ["content_id", "bigint", "", "内容编号", "", ""],
            ["created_at", "timestamp", "", "收藏时间", "", "CURRENT_TIMESTAMP"],
        ],
    },
    {
        "num": 4,
        "name": "活动表（activities）",
        "desc": "储存非遗活动的基本信息，包括时间、地点、人数限制和发布状态",
        "fields": [
            ["id", "bigint", "", "主键", "主键", ""],
            ["title", "varchar", "200", "活动标题", "", ""],
            ["cover_url", "varchar", "255", "封面图地址", "", ""],
            ["location", "varchar", "200", "活动地点", "", ""],
            ["organizer", "varchar", "120", "主办方", "", ""],
            ["start_time", "timestamp", "", "开始时间", "", ""],
            ["end_time", "timestamp", "", "结束时间", "", ""],
            ["max_participants", "int", "", "最大参与人数", "", "0"],
            ["description", "longtext", "4294967295", "活动描述", "", ""],
            ["notes", "text", "", "注意事项", "", ""],
            ["is_featured", "boolean", "", "是否精选", "", "0"],
            ["status", "varchar", "20", "发布状态", "", "draft"],
            ["created_at", "timestamp", "", "创建时间", "", "CURRENT_TIMESTAMP"],
        ],
    },
    {
        "num": 5,
        "name": "活动报名表（activity_registrations）",
        "desc": "储存用户对非遗活动的报名记录和状态",
        "fields": [
            ["id", "bigint", "", "主键", "主键", ""],
            ["activity_id", "bigint", "", "活动编号", "", ""],
            ["user_id", "bigint", "", "用户编号", "", ""],
            ["remark", "varchar", "255", "报名备注", "", ""],
            ["status", "varchar", "20", "报名状态", "", "registered"],
            ["created_at", "timestamp", "", "报名时间", "", "CURRENT_TIMESTAMP"],
            ["updated_at", "timestamp", "", "更新时间", "", "CURRENT_TIMESTAMP"],
        ],
    },
    {
        "num": 6,
        "name": "讨论话题表（discussion_topics）",
        "desc": "储存社区讨论帖子的内容、互动统计和精选标记",
        "fields": [
            ["id", "bigint", "", "主键", "主键", ""],
            ["user_id", "bigint", "", "发帖用户编号", "", ""],
            ["nickname", "varchar", "64", "发帖人昵称", "", ""],
            ["title", "varchar", "200", "话题标题", "", ""],
            ["content", "longtext", "4294967295", "话题内容", "", ""],
            ["cover_url", "varchar", "255", "封面图地址", "", ""],
            ["image_urls", "text", "", "图片地址列表(JSON)", "", ""],
            ["is_featured", "boolean", "", "是否精选", "", "0"],
            ["like_count", "int", "", "点赞数", "", "0"],
            ["favorite_count", "int", "", "收藏数", "", "0"],
            ["comment_count", "int", "", "评论数", "", "0"],
            ["created_at", "timestamp", "", "创建时间", "", "CURRENT_TIMESTAMP"],
        ],
    },
    {
        "num": 7,
        "name": "评论表（discussion_comments）",
        "desc": "储存用户对讨论话题的评论内容",
        "fields": [
            ["id", "bigint", "", "主键", "主键", ""],
            ["topic_id", "bigint", "", "话题编号", "", ""],
            ["user_id", "bigint", "", "评论用户编号", "", ""],
            ["nickname", "text", "", "评论人昵称", "", ""],
            ["content", "longtext", "4294967295", "评论内容", "", ""],
            ["created_at", "timestamp", "", "评论时间", "", "CURRENT_TIMESTAMP"],
        ],
    },
    {
        "num": 8,
        "name": "点赞表（discussion_likes）",
        "desc": "储存用户对讨论话题的点赞记录",
        "fields": [
            ["id", "bigint", "", "主键", "主键", ""],
            ["topic_id", "bigint", "", "话题编号", "", ""],
            ["user_id", "bigint", "", "用户编号", "", ""],
            ["created_at", "timestamp", "", "点赞时间", "", "CURRENT_TIMESTAMP"],
        ],
    },
    {
        "num": 9,
        "name": "话题收藏表（discussion_favorites）",
        "desc": "储存用户对讨论话题的收藏记录",
        "fields": [
            ["id", "bigint", "", "主键", "主键", ""],
            ["topic_id", "bigint", "", "话题编号", "", ""],
            ["user_id", "bigint", "", "用户编号", "", ""],
            ["created_at", "timestamp", "", "收藏时间", "", "CURRENT_TIMESTAMP"],
        ],
    },
    {
        "num": 10,
        "name": "话题标签表（discussion_topic_tags）",
        "desc": "储存讨论话题的标签信息，用于话题分类和检索",
        "fields": [
            ["id", "bigint", "", "主键", "主键", ""],
            ["topic_id", "bigint", "", "话题编号", "", ""],
            ["tag", "varchar", "50", "标签文本", "", ""],
            ["created_at", "timestamp", "", "创建时间", "", "CURRENT_TIMESTAMP"],
        ],
    },
    {
        "num": 11,
        "name": "推荐记录表（recommend_logs）",
        "desc": "储存推荐行为的日志记录，包括行为类型、推荐对象和推荐解释",
        "fields": [
            ["id", "bigint", "", "主键", "主键", ""],
            ["user_id", "bigint", "", "用户编号", "", ""],
            ["action", "varchar", "20", "行为类型", "", ""],
            ["target_type", "varchar", "20", "推荐对象类型", "", ""],
            ["target_id", "bigint", "", "对象编号", "", ""],
            ["source_scene", "varchar", "30", "来源场景", "", ""],
            ["explain_json", "text", "", "推荐解释JSON", "", ""],
            ["created_at", "timestamp", "", "记录时间", "", "CURRENT_TIMESTAMP"],
        ],
    },
    {
        "num": 12,
        "name": "CRS会话状态表（crs_sessions）",
        "desc": "储存CRS对话推荐会话的状态信息，包括当前模式、轮次和已问属性",
        "fields": [
            ["id", "bigint", "", "主键", "主键", ""],
            ["user_id", "bigint", "", "用户编号", "", ""],
            ["session_id", "varchar", "36", "会话UUID", "", ""],
            ["mode", "varchar", "20", "当前模式", "", "cold_start"],
            ["turn_count", "int", "", "对话轮次", "", "0"],
            ["last_ask_attribute", "varchar", "30", "最近ASK属性", "", ""],
            ["last_ask_id", "varchar", "10", "最近ASK模板编号", "", ""],
            ["asked_attributes", "text", "", "已问属性列表(JSON)", "", ""],
            ["context_summary", "text", "", "对话上下文摘要", "", ""],
            ["is_active", "int", "", "会话是否活跃", "", "1"],
            ["created_at", "timestamp", "", "创建时间", "", "CURRENT_TIMESTAMP"],
            ["updated_at", "timestamp", "", "更新时间", "", "CURRENT_TIMESTAMP"],
        ],
    },
    {
        "num": 13,
        "name": "CRS提问记录表（crs_ask_logs）",
        "desc": "储存CRS对话推荐中ASK模板的提问和用户回答记录",
        "fields": [
            ["id", "bigint", "", "主键", "主键", ""],
            ["user_id", "bigint", "", "用户编号", "", ""],
            ["session_id", "varchar", "36", "会话编号", "", ""],
            ["ask_id", "varchar", "10", "ASK模板编号", "", ""],
            ["attribute", "varchar", "30", "提问属性", "", ""],
            ["question_text", "text", "", "提问原文", "", ""],
            ["answer", "text", "", "用户回答", "", ""],
            ["score_delta", "int", "", "置信度增量", "", "0"],
            ["created_at", "timestamp", "", "记录时间", "", "CURRENT_TIMESTAMP"],
        ],
    },
    {
        "num": 14,
        "name": "AI问答记录表（ai_qa_logs）",
        "desc": "储存AI数字人问答的完整记录，包括问题、回答、来源和置信度",
        "fields": [
            ["id", "bigint", "", "主键", "主键", ""],
            ["user_id", "bigint", "", "用户编号", "", ""],
            ["question", "text", "", "用户提问", "", ""],
            ["answer", "longtext", "4294967295", "AI回答", "", ""],
            ["source", "varchar", "30", "回答来源", "", ""],
            ["confidence", "numeric", "4,2", "回答置信度", "", ""],
            ["created_at", "timestamp", "", "记录时间", "", "CURRENT_TIMESTAMP"],
        ],
    },
    {
        "num": 15,
        "name": "本地知识库表（local_knowledge_base）",
        "desc": "储存经过采集和审核的非遗专业知识问答对，供AI数字人检索调用",
        "fields": [
            ["id", "bigint", "", "主键", "主键", ""],
            ["question", "text", "", "知识库问题", "", ""],
            ["answer", "longtext", "4294967295", "知识库回答", "", ""],
            ["qa_answer", "text", "", "问答对答案", "", ""],
            ["keywords", "varchar", "255", "关键词", "", ""],
            ["chapter", "varchar", "120", "所属章节", "", ""],
            ["sub_chapter", "varchar", "120", "子章节", "", ""],
            ["source", "varchar", "100", "来源", "", ""],
            ["status", "varchar", "20", "状态", "", "active"],
            ["created_at", "timestamp", "", "创建时间", "", "CURRENT_TIMESTAMP"],
        ],
    },
    {
        "num": 16,
        "name": "电子资料表（electronic_materials）",
        "desc": "储存平台上传的电子资料文件信息，包括文件地址、类型和摘要",
        "fields": [
            ["id", "bigint", "", "主键", "主键", ""],
            ["title", "varchar", "200", "资料标题", "", ""],
            ["file_url", "varchar", "255", "文件地址", "", ""],
            ["file_type", "varchar", "30", "文件类型", "", ""],
            ["summary", "text", "", "摘要", "", ""],
            ["status", "varchar", "20", "状态", "", "active"],
            ["created_by", "bigint", "", "创建人编号", "", ""],
            ["created_at", "timestamp", "", "创建时间", "", "CURRENT_TIMESTAMP"],
        ],
    },
]


def _set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _make_three_line_table(doc, headers, rows, caption=""):
    if caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap_para.add_run(caption)
        run.font.size = Pt(10.5)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        cap_para.paragraph_format.space_after = Pt(4)
        cap_para.paragraph_format.space_before = Pt(6)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(header))
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        _set_cell_shading(cell, "F2F2F2")

    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            if col_idx in (1, 2, 4, 5):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()
            tc_borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'</w:tcBorders>'
            )
            existing = tc_pr.find(qn("w:tcBorders"))
            if existing is not None:
                tc_pr.remove(existing)
            tc_pr.append(tc_borders)

    for cell in table.rows[0].cells:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
            f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tcBorders>'
        )
        existing = tc_pr.find(qn("w:tcBorders"))
        if existing is not None:
            tc_pr.remove(existing)
        tc_pr.append(tc_borders)

    for cell in table.rows[-1].cells:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tcBorders>'
        )
        existing = tc_pr.find(qn("w:tcBorders"))
        if existing is not None:
            tc_pr.remove(existing)
        tc_pr.append(tc_borders)

    return table


def _add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    return heading


def _add_body_text(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return para


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing = 1.5

    _add_heading(doc, "3.5 系统数据字典", level=2)

    _add_body_text(doc, (
        "储存在数据库中的是数据表，数据间的关联性由这些数据表构成。"
        "系统数据库共包含16张数据表，按照业务领域可以划分为用户域、内容域、活动域、社区域、推荐域和知识域六个部分。"
        "部分系统所涉及的数据表如下所示。"
    ))

    for i, tbl in enumerate(TABLES):
        if i > 0:
            doc.add_paragraph()

        para = doc.add_paragraph()
        run = para.add_run(f"（{tbl['num']}）{tbl['name']}")
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        _add_body_text(doc, f"表结构如下表{tbl['num']}所示：")

        _make_three_line_table(
            doc,
            HEADERS,
            tbl["fields"],
            caption=f"表{tbl['num']} {tbl['name']}",
        )

    output_path = r"d:\桌面\毕业设计\论文正文\3.5系统数据字典.docx"
    doc.save(output_path)
    print(f"Word文档已生成: {output_path}")


if __name__ == "__main__":
    main()
