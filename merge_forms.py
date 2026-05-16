#!/usr/bin/env python3
"""Merge all teaching management forms into one document following the TOC order"""
import os, sys
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from copy import deepcopy
import xml.etree.ElementTree as ET

BASE = r'd:\桌面\毕业设计\论文正文\最终正式版\教学管理材料'
TEMPLATE_BASE = r'd:\桌面\论文模板'
OUTPUT = os.path.join(BASE, '教学管理材料（完整版）.docx')

# Files in order per 附件3 TOC
FILES_ORDER = [
    '附件4. 本科毕业论文（设计）拟题审批表.docx',
    '附件6. 本科毕业论文（设计）任务书.docx',
    '附件7. 本科毕业论文（设计）开题报告.docx',
    '附件8. 本科毕业论文（设计）进展情况记录表.docx',
    '附件9. 本科毕业论文（设计）中期检查表.docx',
    '附件10. 本科毕业论文（设计）答辩记录表.docx',
]

def get_section_width(doc):
    """Get the page width from first section"""
    for section in doc.sections:
        return section.page_width
    return None

def set_page_size(section, width, height):
    """Set page size for a section"""
    section.page_width = width
    section.page_height = height

def copy_paragraph_style(src_para, dst_para):
    """Copy paragraph formatting"""
    dst_para.alignment = src_para.alignment
    if src_para.style:
        try:
            dst_para.style = src_para.style
        except:
            pass

def copy_run_style(src_run, dst_run):
    """Copy run formatting"""
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    if src_run.font.size:
        dst_run.font.size = src_run.font.size
    if src_run.font.name:
        dst_run.font.name = src_run.font.name
    if src_run.font.color and src_run.font.color.rgb:
        dst_run.font.color.rgb = src_run.font.color.rgb

def copy_paragraph(src_para, dst_doc):
    """Copy a paragraph from source to destination document"""
    new_para = dst_doc.add_paragraph()
    copy_paragraph_style(src_para, new_para)

    for src_run in src_para.runs:
        new_run = new_para.add_run(src_run.text)
        copy_run_style(src_run, new_run)

    # Remove the default empty run if we added others
    if len(new_para.runs) > 1 and new_para.runs[0].text == '':
        pass  # Keep structure

    return new_para

def copy_table(src_table, dst_doc):
    """Copy a table from source to destination"""
    rows = len(src_table.rows)
    cols = len(src_table.columns)
    new_table = dst_doc.add_table(rows=rows, cols=cols)

    # Copy table style
    if src_table.style:
        try:
            new_table.style = src_table.style
        except:
            pass

    # Copy table alignment
    new_table.alignment = src_table.alignment

    for ri, src_row in enumerate(src_table.rows):
        dst_row = new_table.rows[ri]
        # Copy row height
        if src_row.height:
            try:
                dst_row.height = src_row.height
            except:
                pass

        for ci, src_cell in enumerate(src_row.cells):
            dst_cell = dst_row.cells[ci]

            # Copy cell width
            if src_cell.width:
                try:
                    dst_cell.width = src_cell.width
                except:
                    pass

            # Copy cell content
            for pi, src_para in enumerate(src_cell.paragraphs):
                if pi == 0:
                    dst_para = dst_cell.paragraphs[0]
                else:
                    dst_para = dst_cell.add_paragraph()

                copy_paragraph_style(src_para, dst_para)

                # Clear default text
                for run in dst_para.runs:
                    run.text = ''

                # Copy runs
                for src_run in src_para.runs:
                    new_run = dst_para.add_run(src_run.text)
                    copy_run_style(src_run, new_run)

                # Remove initial empty run if we added content
                if len(dst_para.runs) > 1:
                    for run in list(dst_para.runs):
                        if run.text == '' and run is not dst_para.runs[-1]:
                            run._element.getparent().remove(run._element)

    return new_table

def add_page_break(doc):
    """Add a page break"""
    para = doc.add_paragraph()
    run = para.add_run()
    run._element.append(ET.fromstring('<w:br xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="page"/>'))
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)

def embed_document(source_path, dest_doc):
    """Embed all content from a source docx into destination doc"""
    print(f"  Embedding: {os.path.basename(source_path)}")
    src = Document(source_path)

    # Copy all elements from source body to destination
    body = src.element.body

    # Process elements in order
    for child in list(body):
        # Handle section properties (page breaks between documents)
        if child.tag == qn('w:sectPr'):
            continue

        if child.tag == qn('w:p'):
            # It's a paragraph
            new_para = dest_doc.add_paragraph()
            # Copy paragraph properties
            pPr = child.find(qn('w:pPr'))
            if pPr is not None:
                new_para._element.get_or_add_pPr().clear()
                for prop_child in list(pPr):
                    new_para._element.get_or_add_pPr().append(deepcopy(prop_child))
            # Copy runs
            for run_elem in child.findall(qn('w:r')):
                new_para._element.append(deepcopy(run_elem))
            # Remove empty default run
            for run in list(new_para.runs):
                if run.text == '' and len(list(new_para.runs)) > 1:
                    run._element.getparent().remove(run._element)

        elif child.tag == qn('w:tbl'):
            # It's a table - we need to add it properly
            new_table_elem = deepcopy(child)
            dest_doc.element.body.append(new_table_elem)

        elif child.tag == qn('w:sdt'):
            # Structured document tag (content controls)
            new_sdt = deepcopy(child)
            dest_doc.element.body.append(new_sdt)

def main():
    print("Creating combined teaching management document...")

    # Start with 附件3 as the base (it has the cover and TOC)
    base_file = os.path.join(BASE, '附件3. 本科毕业论文（设计）教学管理材料封面及目录.docx')
    print(f"Base: {os.path.basename(base_file)}")
    dest = Document(base_file)

    # Get page dimensions from base
    for section in dest.sections:
        page_w = section.page_width
        page_h = section.page_height
        break

    # Add page break and embed each subsequent document
    for filename in FILES_ORDER:
        filepath = os.path.join(BASE, filename)
        if not os.path.exists(filepath):
            print(f"  WARNING: {filename} not found, skipping")
            continue

        # Add a page break before each new section
        add_page_break(dest)

        # Embed the document
        try:
            embed_document(filepath, dest)
        except Exception as e:
            print(f"  ERROR embedding {filename}: {e}")

    # Add final note about 查重报告
    add_page_break(dest)
    note_para = dest.add_paragraph()
    note_run = note_para.add_run('附件：毕业论文（设计）查重报告')
    note_run.bold = True
    note_run.font.size = Pt(14)
    note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    note2 = dest.add_paragraph()
    note2.paragraph_format.space_before = Pt(20)
    note_run2 = note2.add_run('（查重报告需自行通过学校指定知网官方通道完成查重后附入，此处预留位置。）')
    note_run2.font.size = Pt(11)
    note2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    dest.save(OUTPUT)
    print(f"\nDone! Saved to: {OUTPUT}")
    print(f"File size: {os.path.getsize(OUTPUT) / 1024:.1f} KB")

if __name__ == '__main__':
    main()
