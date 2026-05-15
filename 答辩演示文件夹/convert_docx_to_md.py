import sys
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_docx_to_md(docx_path, md_path):
    doc = Document(docx_path)
    lines = []

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        if not text and style_name != "Heading 1":
            lines.append("")
            continue

        if "Heading 1" in style_name:
            lines.append(f"# {text}")
            lines.append("")
        elif "Heading 2" in style_name:
            lines.append(f"## {text}")
            lines.append("")
        elif "Heading 3" in style_name:
            lines.append(f"### {text}")
            lines.append("")
        elif "Heading 4" in style_name:
            lines.append(f"#### {text}")
            lines.append("")
        elif "Heading 5" in style_name:
            lines.append(f"##### {text}")
            lines.append("")
        elif "Title" in style_name:
            lines.append(f"# {text}")
            lines.append("")
        elif "Caption" in style_name:
            lines.append(f"**{text}**")
            lines.append("")
        else:
            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                lines.append(f"<div align=\"center\">{text}</div>")
                lines.append("")
            else:
                is_bold_all = all(
                    run.bold for run in para.runs if run.text.strip()
                ) if para.runs else False
                if is_bold_all and len(para.runs) > 0:
                    lines.append(f"**{text}**")
                    lines.append("")
                else:
                    lines.append(text)
                    lines.append("")

    for table in doc.tables:
        lines.append("")
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            line = "| " + " | ".join(cells) + " |"
            lines.append(line)
            if i == 0:
                sep = "| " + " | ".join(["---"] * len(cells)) + " |"
                lines.append(sep)
        lines.append("")

    md_content = "\n".join(lines)

    while "\n\n\n\n" in md_content:
        md_content = md_content.replace("\n\n\n\n", "\n\n\n")

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)

    print(f"转换完成: {md_path}")
    print(f"文件大小: {os.path.getsize(md_path)} 字节")
    print(f"总行数: {md_content.count(chr(10)) + 1}")

if __name__ == "__main__":
    docx_path = r"d:\桌面\毕业设计\大数据222-202210274225-王子轩-基于CRS推荐与AI数字人的非遗文化传播系统实现.docx"
    md_path = r"d:\桌面\毕业设计\答辩演示文件夹\论文全文.md"
    convert_docx_to_md(docx_path, md_path)
